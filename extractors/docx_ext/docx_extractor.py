
from pathlib import Path
from docx import Document
import zipfile
import shutil

from extractors.docx_ext.docx_ocr import run_ocr_on_images


# mergeCode 루트 기준 output 폴더.
# DOCX 추출 결과는 output/text, 내부 이미지는 output/docx_images 아래에 저장한다.
output_dir = Path(__file__).resolve().parents[2] / "output"


def load_docx(docx_path: Path):
    """python-docx로 DOCX 문서를 연다."""
    return Document(docx_path)


def extract_paragraphs(document: Document):
    """DOCX 본문에 있는 일반 문단 텍스트를 순서대로 추출한다."""
    texts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            texts.append(text)

    return texts


def extract_tables(document: Document):
    """Word 표 안의 텍스트를 행 단위로 추출한다."""
    texts = []

    for table_index, table in enumerate(document.tables, start=1):
        texts.append(f"\n===== TABLE {table_index} START =====")

        for row in table.rows:
            row_texts = []

            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " ")
                row_texts.append(cell_text)

            texts.append(" | ".join(row_texts))

        texts.append(f"===== TABLE {table_index} END =====\n")

    return texts


def extract_images_from_docx(docx_path: Path, image_output_dir: Path):
    """DOCX 압축 파일 내부의 word/media 이미지들을 별도 폴더로 복사한다."""
    image_output_dir.mkdir(parents=True, exist_ok=True)

    extracted_images = []

    with zipfile.ZipFile(docx_path, "r") as docx_zip:
        for file_name in docx_zip.namelist():
            if file_name.startswith("word/media/"):
                image_name = Path(file_name).name
                image_path = image_output_dir / image_name

                with docx_zip.open(file_name) as source:
                    with open(image_path, "wb") as target:
                        shutil.copyfileobj(source, target)

                extracted_images.append(image_path)

    return extracted_images


def docx_extractor(file_path: Path):
    """DOCX 텍스트 직접 추출 + 내부 이미지 OCR을 합친 하이브리드 추출기."""
    document = load_docx(file_path)

    # 1) 일반 문단과 표는 OCR보다 직접 추출이 정확하므로 먼저 가져온다.
    texts = []
    texts.extend(extract_paragraphs(document))
    texts.extend(extract_tables(document))

    # 2) DOCX 내부에 삽입된 이미지는 word/media에서 꺼낸 뒤 OCR 대상으로 넘긴다.
    image_paths = extract_images_from_docx(
        file_path,
        output_dir / "docx_images" / file_path.stem
    )

    # 3) OCR 입력용 이미지는 docx_ocr.py에서 2배 RGB 업스케일 후 처리한다.
    image_ocr_texts = run_ocr_on_images(
        image_paths,
        output_dir / "docx_images_upscaled" / file_path.stem,
    )
    texts.extend(image_ocr_texts)

    # 4) 직접 추출 텍스트와 이미지 OCR 텍스트를 하나의 txt 결과로 저장한다.
    result_text = "\n".join(texts)

    result_path = output_dir / "text" / f"{file_path.stem}_hybrid_extract.txt"
    result_path.parent.mkdir(parents=True, exist_ok=True)
    result_path.write_text(result_text, encoding="utf-8")

    # 이후 chunk/json/RAG 단계에서 재사용할 수 있도록 주요 산출물을 dict로 반환한다.
    return {
        "file_name": file_path.name,
        "text": result_text,
        "image_paths": image_paths,
        "ocr_texts": image_ocr_texts,
        "text_path": result_path,
    }
